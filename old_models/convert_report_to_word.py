import re
import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Thư viện 'python-docx' chưa được cài đặt.")
    print("👉 Vui lòng chạy lệnh: pipe install python-docx")
    sys.exit(1)

def clean_text(text):
    """Remove markdown syntax from text for basic display."""
    # Remove Bold **text** -> text (we handle formatting separately in a real parser, 
    # but efficient simple adding runs is better)
    return text.strip()

def add_formatted_paragraph(paragraph, text):
    """Parses bold (**), italic (*) and code (`) and adds runs."""
    # Split by bold first: **text**
    # This is a simple regex parser, supports non-nested formatting
    
    # Regex for **bold**, *italic*, `code`
    # Tokenize: Split into chunks of (text, type)
    tokens = []
    
    # Simple state machine or regex split
    # Pattern: capture **...**, *...*, `...`
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(200, 0, 0) # Dark Red for code
        else:
            paragraph.add_run(part)

def convert_md_to_docx(md_file, docx_file):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []
    
    for line in lines:
        line = line.strip()
        
        # --- TABLE HANDLING ---
        if line.startswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                # Process table
                # Filter separator lines like |---|---|
                data_rows = [r for r in table_lines if not set(r.replace('|', '').strip()) <= {'-'}]
                if data_rows:
                    # Create table
                    # Count cols
                    header = [c.strip() for c in data_rows[0].strip('|').split('|')]
                    table = doc.add_table(rows=len(data_rows), cols=len(header))
                    table.style = 'Table Grid'
                    
                    for r_idx, row_str in enumerate(data_rows):
                        cells = [c.strip() for c in row_str.strip('|').split('|')]
                        row_cells = table.rows[r_idx].cells
                        for c_idx, cell_text in enumerate(cells):
                            if c_idx < len(row_cells):
                                # cell_text handles bold? simple text for now
                                row_cells[c_idx].text = cell_text
                
                in_table = False
                table_lines = []
        
        if not line:
            continue

        # --- HEADERS ---
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            
        # --- IMAGES ---
        elif line.startswith('!['):
            # ![Alt](path)
            match = re.search(r'\!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt = match.group(1)
                path = match.group(2)
                if os.path.exists(path):
                    try:
                        doc.add_picture(path, width=Inches(6.0))
                        # Add caption
                        caption = doc.add_paragraph(alt)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.style = 'Caption'
                    except Exception as e:
                        doc.add_paragraph(f"[Image Error: {e}]")
                else:
                    doc.add_paragraph(f"[Image not found: {path}]", style='Emphasis')

        # --- BLOCKQUOTES ---
        elif line.startswith('> '):
            p = doc.add_paragraph()
            add_formatted_paragraph(p, line[2:])
            p.style = 'Intense Quote'
            
        # --- LISTS ---
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_paragraph(p, line[2:])
            
        elif re.match(r'^\d+\.\s', line):
            # Ordered list
            text_part = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_paragraph(p, text_part)
            
        # --- HORIZONTAL RULE ---
        elif line.startswith('---'):
            doc.add_paragraph('_' * 40, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # --- NORMAL TEXT ---
        else:
            p = doc.add_paragraph()
            add_formatted_paragraph(p, line)

    doc.save(docx_file)
    print(f"✅ Successfully converted '{md_file}' to '{docx_file}'")

if __name__ == "__main__":
    md_path = 'REPORT_CONTENT.md'
    docx_path = 'Final_Project_Report.docx'
    
    if not os.path.exists(md_path):
        print(f"File {md_path} not found.")
    else:
        convert_md_to_docx(md_path, docx_path)
